from langchain_core.output_parsers import StrOutputParser
from PyPDF2 import PdfReader
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.output_parsers import StrOutputParser
from docx import Document
# We need to import the specific object types to check against
from docx.document import Document as DocumentObject
from docx.table import Table

class Reader:
    def __init__(self):
        pass
    
    @staticmethod
    def get_reader(file):
        if file.endswith(".pdf"):
            return PDFReader()
        elif file.endswith(".docx"):
            return DocxReader()
        else:
            raise ValueError("Unsupported file type")

    def read(self, file_path: str) -> str:
        pass
    

class PDFReader(Reader):

    def read(self, file_path: str) -> str:
        reader = PdfReader(file_path)
        text = ""
        for page in (reader.pages):
            text += page.extract_text()
        return text
    
class DocxReader(Reader):

    def read(self, file_path: str) -> str:
        """
        Reads paragraphs and tables from a .docx file and prints them in sequential order.
        """
        try:
            doc = Document(file_path)
            full_content = ""

            # Pointers to track our position in the paragraphs and tables lists
            para_idx, table_idx = 0, 0

            # Iterate through the direct children of the document's body
            for block in doc.element.body:
                # Check if the block is a paragraph element
                if block.tag.endswith('p'):
                    paragraph = doc.paragraphs[para_idx]
                    full_content += paragraph.text
                    para_idx += 1
                # Check if the block is a table element
                elif block.tag.endswith('tbl'):
                    table = doc.tables[table_idx]
                    # Iterate through table rows and cells
                    for row in table.rows:
                        row_text = "\t|\t".join(cell.text for cell in row.cells)
                        full_content += f"\t{row_text}"
                    table_idx += 1
            return full_content
        except FileNotFoundError:
            print(f"Error: The file '{file_path}' was not found.")
        except Exception as e:
            print(f"An error occurred: {e}")
    
# class excelReader(Reader):
#     def __init__(self):
#         pass

#     def read(self, file_path: str) -> str:
#         df = pd.read_excel(file_path)
#         text = df.to_string()
#         return text